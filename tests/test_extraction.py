"""
tests/test_extraction.py

Tests for: app/services/extraction.py
Covers: _normalise(), _read() for txt/docx/pdf, extract_text() pipeline,
        unsupported formats, unicode handling, empty files.
"""

import os
import pytest
import tempfile

from app.services.extraction import extract_text, _normalise


# ---------------------------------------------------------------------------
# Unit: _normalise
# ---------------------------------------------------------------------------

class TestNormalise:
    def test_lowercases(self):
        assert _normalise("Hello WORLD") == "hello world"

    def test_strips_punctuation(self):
        result = _normalise("Hello, world! It's great.")
        assert "," not in result
        assert "!" not in result
        assert "'" not in result

    def test_collapses_whitespace(self):
        assert "  " not in _normalise("too   many    spaces")

    def test_strips_leading_trailing_whitespace(self):
        assert _normalise("  hello  ") == "hello"

    def test_empty_string_returns_empty(self):
        assert _normalise("") == ""

    def test_unicode_normalised_to_ascii(self):
        # accented characters should survive or be dropped, not crash
        result = _normalise("café résumé naïve")
        assert isinstance(result, str)
        # accents stripped via NFKD + ascii encode
        assert "é" not in result

    def test_numbers_preserved(self):
        assert "42" in _normalise("There are 42 items.")

    def test_only_punctuation_returns_empty(self):
        assert _normalise("!!! ???") == ""


# ---------------------------------------------------------------------------
# Integration: extract_text from real files
# ---------------------------------------------------------------------------

class TestExtractTxt:
    def test_reads_txt_content(self, tmp_path):
        f = tmp_path / "essay.txt"
        f.write_text("Hello World This Is A Test")
        result = extract_text(str(f))
        assert "hello" in result
        assert "world" in result

    def test_normalisation_applied_to_txt(self, tmp_path):
        f = tmp_path / "essay.txt"
        f.write_text("UPPERCASE AND punctuation!!!")
        result = extract_text(str(f))
        assert result == result.lower()
        assert "!" not in result

    def test_empty_txt_returns_empty_string(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("")
        assert extract_text(str(f)) == ""

    def test_utf8_encoding_handled(self, tmp_path):
        f = tmp_path / "essay.txt"
        f.write_text("Content with unicode: café", encoding="utf-8")
        result = extract_text(str(f))
        assert isinstance(result, str)


class TestExtractDocx:
    def test_reads_docx_paragraphs(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        path = str(tmp_path / "essay.docx")
        doc = Document()
        doc.add_paragraph("The mitochondria is the powerhouse of the cell.")
        doc.save(path)
        result = extract_text(path)
        assert "mitochondria" in result

    def test_multipage_docx_concatenated(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        path = str(tmp_path / "multi.docx")
        doc = Document()
        for i in range(5):
            doc.add_paragraph(f"Paragraph number {i} with content.")
        doc.save(path)
        result = extract_text(path)
        assert "paragraph" in result


class TestExtractPdf:
    def test_reads_pdf_text(self, tmp_path):
        pytest.importorskip("pypdf")
        # Use pypdf to create a minimal PDF for testing
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        path = str(tmp_path / "essay.pdf")
        with open(path, "wb") as f:
            writer.write(f)
        # Blank page — just assert no crash and returns string
        result = extract_text(path)
        assert isinstance(result, str)


class TestExtractUnsupported:
    def test_unsupported_extension_raises_value_error(self, tmp_path):
        f = tmp_path / "essay.csv"
        f.write_text("col1,col2")
        with pytest.raises(ValueError, match="Unsupported"):
            extract_text(str(f))

    def test_unknown_extension_raises_value_error(self, tmp_path):
        f = tmp_path / "mystery.xyz"
        f.write_text("data")
        with pytest.raises(ValueError):
            extract_text(str(f))